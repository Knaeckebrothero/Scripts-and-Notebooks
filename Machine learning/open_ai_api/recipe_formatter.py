import os
import openai
from dotenv import load_dotenv
from dev_tools import read as rd
from docx import Document


# Template for the recipes
def create_recipe(title, portions, kcal, main_ingredients, spices, extras,
                  preparation_time, preparation_steps, cooking_steps, serving_steps
                  ) -> Document():
    # Create a new document
    document = Document()

    # Title
    document.add_heading(title, level=1)

    # Ingredients
    document.add_heading('Zutaten', level=2)
    document.add_paragraph(f'Für {portions} Portionen ca. {kcal} kcal')
    document.add_heading('Hauptzutaten', level=3)
    for ingredient in main_ingredients:
        document.add_paragraph(ingredient, style='ListBullet')
    document.add_heading('Gewürze', level=3)
    for spice in spices:
        document.add_paragraph(spice, style='ListBullet')
    document.add_heading('Weiteres', level=3)
    for extra in extras:
        document.add_paragraph(extra, style='ListBullet')

    # Preparation
    document.add_heading('Zubereitung', level=2)
    document.add_paragraph(f'Gesamtzeit ca. {preparation_time} Minuten')
    document.add_heading('Vorbereitung', level=3)
    for step in preparation_steps:
        document.add_paragraph(step, style='ListBullet')
    document.add_heading('Zubereitung', level=3)
    for step in cooking_steps:
        document.add_paragraph(step, style='ListBullet')
    document.add_heading('Anrichten', level=3)
    for step in serving_steps:
        document.add_paragraph(step, style='ListBullet')

    # Return the document .save(f'{title}.docx')
    return document


def function_call(content, description):
    # Define the function as an array
    functions = [description]

    # Define the messages as an array
    messages = [
        {"role": "system", "content": "You are a helpful assistant."},
        {"role": "user", "content": content},
        # ... additional messages if needed
    ]

    # Make the API call
    api_response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo-16k-0613",
        messages=messages,
        functions=functions,
        function_call="extract_recipe_details"
    )
    return api_response


# Load the API key from the .env file.
load_dotenv()
openai.api_key = os.getenv('OPENAI_KEY')

# Make the API call
response = function_call(rd.text_data('recipe.txt'), rd.json_data('recipe_description.json'))

# Extract the details from the API response
recipe_details = response['choices'][0]['message']['content']['extract_recipe_details']

# Create the recipe
document = create_recipe(
    title=recipe_details['title'],
    portions=recipe_details['portions'],
    kcal=recipe_details['kcal'],
    main_ingredients=recipe_details['main_ingredients'],
    spices=recipe_details['spices'],
    extras=recipe_details['extras'],
    preparation_time=recipe_details['preparation_time'],
    preparation_steps=recipe_details['preparation_steps'],
    cooking_steps=recipe_details['cooking_steps'],
    serving_steps=recipe_details['serving_steps']
)

# Save the document
document.save(f'Machine learning/open_ai_api/{recipe_details["title"]}.docx')

print('Done!')
